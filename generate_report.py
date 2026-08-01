import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_report():
    doc = docx.Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Heading styles
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = 'Arial'
        h_font.color.rgb = RGBColor(0, 51, 153)
        h_font.bold = True
    
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(12)

    def add_title(text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER, color=None, bold=True):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
            
        if color:
            run.font.color.rgb = color
        else:
            run.font.color.rgb = RGBColor(0, 51, 153) # Default dark blue
        return p
        
    def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    # --- Title Page ---
    # Removed invalid picture addition
    # Placeholder for logo
    # Actually, we don't have the logo image file, so we'll just add text. I'll comment out image insertion.
    
    # Let's just use text for everything to ensure it runs
    add_title("DEEPFAKE DETECTION USING CONVOLUTIONAL NEURAL NETWORKS\n", level=1)
    add_title("MAJOR PROJECT REPORT\n", level=2, color=RGBColor(0, 0, 0))
    
    add_title("This report is submitted in partial fulfillment of the requirements for the degree of\nMaster of Computer Applications by\n", level=3, color=RGBColor(0, 0, 0), bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_title("JOYGANESH BARAT [120710242004]\n&\nSAYAN BARAT [120710242009]\n", level=2, color=RGBColor(0, 0, 0))
    
    add_title("UNDER THE GUIDANCE OF\nPROF. ANUPAM BAIDYA, ASSISTANT PROFESSOR,\nDEPARTMENT OF COMPUTER APPLICATION, BCREC.\n", level=3, color=RGBColor(0, 0, 0))
    
    add_title("DR. B. C. ROY ENGINEERING COLLEGE\nFuljhore, Jemua Road, Durgapur – 713206\nWest Bengal, India\n2025 - 2026", level=3, color=RGBColor(0, 0, 0), bold=False)
    
    doc.add_page_break()

    # --- Certificate ---
    add_title("CERTIFICATE")
    add_paragraph("This is to certify that Joyganesh Barat (Roll No: 120710242004) and Sayan Barat (Roll No: 120710242009), students of the Department of Computer Applications, BCREC have successfully completed their Major Project entitled “Deepfake Detection Using Convolutional Neural Networks” during the academic session 2025–2026.")
    add_paragraph("They have carried out the project work sincerely, diligently, and enthusiastically under my guidance and supervision. To the best of my knowledge and belief, this project work is original, technically sound, and has been completed in accordance with the prescribed academic requirements.")
    add_paragraph("I hereby recommend that the project report submitted by them be accepted as partial fulfilment of the requirements for the award of the degree of Master of Computer Applications (MCA) from Dr. B. C. Roy Engineering College, Durgapur.\n\n\n")
    
    p = doc.add_paragraph()
    p.add_run("Mentor:\t\t\t\t\t\t\tForwarded by:\n")
    p.add_run("________________________\t\t\t\t________________________\n")
    p.add_run("Prof. Anupam Baidya\t\t\t\tProf. (Dr.) Pabitra Kumar Dey\n")
    p.add_run("Assistant Professor\t\t\t\t\tHead\n")
    p.add_run("Department of Computer Applications\t\tDepartment of Computer Applications\n")
    p.add_run("Dr. B. C. Roy Engineering College\t\tDr. B. C. Roy Engineering College")
    
    doc.add_page_break()

    # --- Declaration ---
    add_title("DECLARATION")
    add_paragraph("We, the undersigned students of the Department of Computer Applications (MCA) at Dr. B.C. Roy Engineering College, hereby declares that the project work titled \"Deepfake Detection Using Convolutional Neural Networks\" is a record of original work done by us.")
    add_paragraph("This project aims to address the critical challenge of deepfake detection in modern digital media by automating the detection and classification of manipulated videos. The developed system leverages state-of-the-art Convolutional Neural Networks (CNN), specifically EfficientNet-B4, enhanced with transfer learning to achieve robust binary classification.")
    add_paragraph("We certify that this application was developed as part of our academic curriculum, utilizing our skills in Artificial Intelligence, Computer Vision, and Web Development. We have adhered to ethical coding practices and academic guidelines throughout the development process.")
    add_paragraph("We further declare that the results embodied in this project have not been submitted to any other university or institute for the award of any degree or diploma. We take full responsibility for the authenticity and integrity of the code and documentation provided herein.\n\n\n")
    
    add_paragraph("_________________________\nJoyganesh Barat [120710242004]\n\n\n_________________________\nSayan Barat [120710242009]")
    doc.add_page_break()

    # --- Acknowledgement ---
    add_title("ACKNOWLEDGEMENT")
    add_paragraph("The successful completion of any significant technical project requires the guidance and support of many individuals. We would like to express our sincere gratitude to everyone who contributed to our project, \"Deepfake Detection Using Convolutional Neural Networks.\" Working on this system has been an immense source of knowledge and practical experience in the field of Artificial Intelligence and Computer Vision.")
    add_paragraph("First and foremost, we express our deepest appreciation to our project supervisor, Prof. Anupam Baidya, for his invaluable mentorship, constant encouragement, and expert supervision. His insightful feedback and constructive suggestions were instrumental in shaping the technical architecture and successful implementation of this project.")
    add_paragraph("We are profoundly grateful to Prof. (Dr.) Pabitra Kumar Dey, Head of the Department of Computer Applications, for providing the necessary academic resources and a conducive environment to pursue this research. We also extend our thanks to all faculty members of the Department of Computer Applications for their continued support.")
    add_paragraph("Finally, we thank our peers and friends who provided valuable suggestions, testing feedback, and moral support during the challenging phases of development.\n\n\n")
    add_paragraph("_________________________\nJoyganesh Barat [120710242004]\n\n\n_________________________\nSayan Barat [120710242009]")
    doc.add_page_break()

    # --- Abstract ---
    add_title("ABSTRACT")
    add_paragraph("The proposed major project aims to design and develop a Convolutional Neural Network (CNN) based system capable of automatically detecting and classifying deepfake videos and images with high accuracy. Deepfakes are synthetic media created using Generative Adversarial Networks (GANs) and autoencoders that manipulate facial features, expressions, and voices to create hyper-realistic forgeries, posing severe threats including misinformation, identity theft, financial fraud, political manipulation, and privacy violations.")
    add_paragraph("The model analyzes visual content frame-by-frame to identify subtle inconsistencies and manipulation artifacts such as facial warping, inconsistent lighting patterns, unnatural blinking behaviors, lip-sync mismatches, pixel-level anomalies, and texture irregularities invisible to the human eye. Unlike traditional manual forensic analysis or rule-based detection methods which are time-consuming and inadequate for modern digital content scale, this system employs deep learning-based automated classification achieving binary output (Real vs. Fake) with confidence scores.")
    add_paragraph("The system processes datasets including FaceForensics++, Celeb-DF, and DFDC, performing comprehensive preprocessing with face extraction, frame sampling, normalization, and data augmentation. The core architecture utilizes EfficientNet-B4 pre-trained on ImageNet, enhanced with transfer learning. A minimalist Flask web interface allows users to easily upload videos, view real-time results, frame-by-frame analysis, and Grad-CAM visualizations. The solution operates fully offline without reliance on external APIs, ensuring privacy and robust performance.")
    doc.add_page_break()

    # --- Table of Contents ---
    add_title("TABLE OF CONTENTS")
    toc_items = [
        "Introduction",
        "College Profile",
        "Limitations of Existing System",
        "System Analysis",
        "Objectives",
        "PERT Chart & Gantt Chart",
        "Hardware and Software Requirements",
        "System Design and Development",
        "System Testing and Implementations",
        "Conclusion",
        "Future Scope",
        "Sample Codes",
        "Outcomes",
        "Bibliography"
    ]
    for idx, item in enumerate(toc_items, 1):
        add_paragraph(f"{idx}. {item}.......................................................................( {idx+6:02d} )")
    doc.add_page_break()

    # --- Sections ---
    
    # 1. Introduction
    doc.add_heading("1. INTRODUCTION", level=1)
    doc.add_heading("Project Overview", level=2)
    add_paragraph("In the digital age, deepfakes—AI-generated synthetic media where a person's likeness is replaced or altered—have become a critical challenge. Created using advanced deep learning techniques like GANs and autoencoders, these manipulations are increasingly difficult to distinguish from authentic content. Accessible deepfake generation tools have enabled widespread misuse including political disinformation campaigns, non-consensual explicit content, celebrity impersonation, financial fraud through voice cloning, evidence tampering, and sophisticated social engineering attacks.")
    add_paragraph("The need for automated detection tools is urgent to combat the spread of misinformation, protect individual privacy, and restore trust in digital media. Social media platforms host billions of daily uploads, making manual verification impossible. Convolutional Neural Networks (CNNs) are particularly well-suited for this task due to their ability to extract complex spatial features and detect subtle artifacts within video frames and images.")
    add_paragraph("This project leverages CNN capabilities, specifically EfficientNet-B4 with transfer learning, to build a robust detection tool that assists social media platforms, news organizations, law enforcement agencies, cybersecurity professionals, and general users in verifying the authenticity of digital media, saving time and maintaining information integrity in the digital ecosystem. The project includes a user-friendly Flask-based web interface to provide seamless drag-and-drop video uploads, instantaneous deepfake classification, and explainable AI features such as Grad-CAM heatmap visualizations to highlight manipulated regions.")
    doc.add_page_break()

    # 2. College Profile
    doc.add_heading("2. COLLEGE PROFILE", level=1)
    add_paragraph("Dr. B. C. Roy Engineering College (BCREC) is a prominent autonomous private institution located in Durgapur, West Bengal, offering a wide range of undergraduate and postgraduate programs in engineering, management, computer applications, and pharmacy. Established in 2000, the college is affiliated with the Maulana Abul Kalam Azad University of Technology (MAKAUT) and is approved by the All India Council for Technical Education (AICTE). The college operates as a self-financing institution under the B. C. Roy Society.")
    doc.add_heading("Academic Programs Offered", level=2)
    p = doc.add_paragraph()
    p.add_run("• Bachelor of Technology (B.Tech)\n• Master of Technology (M.Tech)\n• Master of Business Administration (MBA)\n• Master of Computer Applications (MCA)\n• Bachelor of Pharmacy (B.Pharma)")
    doc.add_heading("Accreditations and Affiliations", level=2)
    p = doc.add_paragraph()
    p.add_run("• Approved by AICTE, New Delhi\n• Accredited by NAAC with ‘B+’ Grade\n• NBA accreditation for multiple engineering programs\n• Affiliated to MAKAUT, West Bengal")
    add_paragraph("The college is supported by a dedicated Training and Placement Cell operating from Durgapur and Kolkata, which promotes industry interaction and enhances student employability. With strong academic infrastructure and experienced faculty, BCREC provides a conducive environment for advanced technical projects such as “Deepfake Detection Using Convolutional Neural Networks.”")
    doc.add_page_break()

    # 3. Limitations of Existing System
    doc.add_heading("3. LIMITATIONS OF EXISTING SYSTEM", level=1)
    add_paragraph("Traditional media verification and early-generation deepfake detection approaches face several significant limitations that impact their effectiveness:")
    add_paragraph("Manual Verification Time Consumption: Forensic experts typically spend hours analyzing a single video for inconsistencies, a process that cannot scale given the massive volume of daily uploads to social media and news platforms.")
    add_paragraph("Inadequate Rule-Based Detection: Early automated systems relied on basic rule-based features or simple biometric tracking (like blink rate detection). These are easily bypassed by modern deepfake generators that have learned to accurately synthesize blinks, breathing, and other natural physiological signs.")
    add_paragraph("Lack of Explainability: Many existing AI models act as a \"black box\" providing a simple true/false prediction without context. This lack of interpretability causes distrust among journalists and law enforcement who require evidence of why a video was flagged.")
    add_paragraph("Cloud Dependency and Privacy Risks: Numerous modern deepfake detection tools operate as cloud APIs, requiring users to upload potentially sensitive videos to third-party servers, compromising data privacy and imposing recurring API costs.")
    doc.add_page_break()
    
    # 4. System Analysis
    doc.add_heading("4. SYSTEM ANALYSIS", level=1)
    add_paragraph("System analysis represents a critical phase providing a comprehensive understanding of system requirements, functionality, and constraints. For this project, the analysis focused on creating an effective, privacy-focused application that leverages cutting-edge computer vision to automatically flag manipulated videos.")
    doc.add_heading("Technology Stack Rationale:", level=2)
    p = doc.add_paragraph()
    p.add_run("• Python: Selected as the primary language due to its extensive computer vision and deep learning ecosystem.\n")
    p.add_run("• PyTorch & torchvision: Chosen as the core deep learning framework for its dynamic computational graph, ease of debugging, and comprehensive pre-trained model library (EfficientNet-B4).\n")
    p.add_run("• OpenCV (cv2): Utilized for efficient video frame extraction, manipulation, and face detection preprocessing.\n")
    p.add_run("• Flask: Adopted for developing a lightweight, rapid, and fully local web application that serves the model without heavy dependencies.\n")
    p.add_run("• Grad-CAM: Employed to provide model explainability by generating heatmaps highlighting regions the CNN focused on to make its decision.")
    
    doc.add_heading("System Architecture Overview:", level=2)
    p = doc.add_paragraph()
    p.add_run("The system follows a modular pipeline architecture:\n")
    p.add_run("1. Presentation Layer: Flask-based web interface handling user video uploads and rendering dynamic result dashboards.\n")
    p.add_run("2. Preprocessing Layer: OpenCV-based module for extracting frames and cropping faces using MTCNN or Haar Cascades.\n")
    p.add_run("3. AI Processing Layer: PyTorch EfficientNet-B4 model with a custom classification head evaluating frames and aggregating a final video-level probability score.\n")
    doc.add_page_break()

    # 5. Objectives
    doc.add_heading("5. OBJECTIVES", level=1)
    p = doc.add_paragraph()
    p.add_run("• To design a CNN model capable of accurately classifying images and videos as real or deepfake with high detection accuracy.\n")
    p.add_run("• To create a robust preprocessing pipeline that extracts frames, detects faces, and prepares data from input videos and images.\n")
    p.add_run("• To train the model on diverse benchmark datasets ensuring adaptability to different manipulation techniques.\n")
    p.add_run("• To minimize reliance on third-party APIs by implementing a locally executable system with minimal external dependencies.\n")
    p.add_run("• To evaluate model performance using comprehensive metrics including accuracy, precision, recall, F1-score, and ROC-AUC.\n")
    p.add_run("• To integrate the system into a user-friendly Flask web interface for real-time video upload and evaluation.\n")
    p.add_run("• To provide explainability features through Grad-CAM visualizations highlighting manipulated facial regions.")
    doc.add_page_break()

    # 6. PERT Chart & Gantt Chart
    doc.add_heading("6. PERT CHART & GANTT CHART", level=1)
    add_paragraph("The development of the Deepfake Detection project followed a structured timeline, divided into logical phases to ensure systematic progress and timely completion.")
    
    doc.add_heading("Project Phases", level=2)
    p = doc.add_paragraph()
    p.add_run("• Phase 1: Initiation and Requirement Analysis (August)\n")
    p.add_run("• Phase 2: Dataset Collection and Preprocessing (September)\n")
    p.add_run("• Phase 3: AI Model Architecture and Training (October)\n")
    p.add_run("• Phase 4: Web Interface Development and Integration (November)\n")
    p.add_run("• Phase 5: System Testing, Evaluation, and Documentation (December)")
    add_paragraph("[Please insert your PERT Chart image here.]", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph("[Please insert your Gantt Chart image here.]", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_page_break()

    # 7. Hardware and Software Requirements
    doc.add_heading("7. HARDWARE AND SOFTWARE REQUIREMENTS", level=1)
    doc.add_heading("Hardware Specifications", level=2)
    p = doc.add_paragraph()
    p.add_run("• Processor: Intel Core i5/i7 (10th Gen or newer) or AMD Ryzen equivalent\n")
    p.add_run("• Memory: 16 GB DDR4 RAM minimum (32 GB recommended for large dataset handling)\n")
    p.add_run("• Storage: 512 GB NVMe SSD for fast data throughput\n")
    p.add_run("• Graphics: NVIDIA GPU with CUDA support (e.g., RTX 3060 or better) strongly recommended for model training and rapid inference.")
    
    doc.add_heading("Software Stack", level=2)
    p = doc.add_paragraph()
    p.add_run("• Operating System: Windows 11 Professional / Ubuntu 20.04+\n")
    p.add_run("• Python Distribution: Python 3.8+\n")
    p.add_run("• AI/ML Frameworks: PyTorch 2.0+, torchvision\n")
    p.add_run("• Computer Vision: OpenCV (cv2), face_recognition\n")
    p.add_run("• Web Framework: Flask\n")
    p.add_run("• Data Processing & Visualization: NumPy, pandas, scikit-learn, matplotlib, seaborn, pytorch-grad-cam")
    doc.add_page_break()

    # 8. System Design and Development
    doc.add_heading("8. SYSTEM DESIGN AND DEVELOPMENT", level=1)
    add_paragraph("The development phase translates theoretical models into functional software. The Deepfake Detection project focuses on creating a high-accuracy classification pipeline utilizing transfer learning.")
    
    doc.add_heading("Model Architecture", level=2)
    add_paragraph("The core architecture relies on EfficientNet-B4 pre-trained on ImageNet. EfficientNet balances network depth, width, and resolution to achieve state-of-the-art accuracy with fewer parameters.")
    p = doc.add_paragraph()
    p.add_run("• Feature Extraction: The convolutional backbone extracts 1792-dimensional feature vectors from 224x224 RGB face crops.\n")
    p.add_run("• Custom Classification Head: The extracted features pass through a custom top layers sequence: Dropout(0.3) -> Linear(1792 to 512) -> BatchNorm1d -> ReLU -> Dropout(0.2) -> Linear(512 to 1) -> Sigmoid activation.\n")
    p.add_run("• Output: A probability score where values closer to 1.0 indicate a 'Fake' and values closer to 0 indicate 'Real'.")
    
    doc.add_heading("Training Pipeline", level=2)
    add_paragraph("The training procedure was divided into two phases to maximize transfer learning efficacy without destroying pre-trained weights:")
    p = doc.add_paragraph()
    p.add_run("• Phase 1 (Head Training): Backbone weights were frozen, and only the custom classification head was trained for 5 epochs with a higher learning rate.\n")
    p.add_run("• Phase 2 (Fine-tuning): The entire network was unfrozen, and training continued with a reduced learning rate and a ReduceLROnPlateau scheduler. Early stopping and gradient clipping were applied to prevent overfitting.")
    add_paragraph("Data augmentation techniques like random horizontal flips, rotation (±15°), color jitter, gaussian blur (to simulate compression), and random erasing were critical to generalize the model against unseen deepfakes.")
    
    doc.add_heading("Web Interface Module", level=2)
    add_paragraph("A minimalist dark theme Flask application provides the user interface. It allows drag-and-drop video uploading. Upon upload, the backend extracts frames, processes them through the PyTorch model, generates a frame-by-frame analysis bar chart, highlights the most suspicious frame, and overlays a Grad-CAM heatmap to show areas of manipulation.")
    doc.add_page_break()

    # 9. System Testing and Implementations
    doc.add_heading("9. SYSTEM TESTING AND IMPLEMENTATION", level=1)
    add_paragraph("System testing ensures the reliability and accuracy of the model across various datasets and conditions.")
    doc.add_heading("Evaluation Metrics", level=2)
    p = doc.add_paragraph()
    p.add_run("• Accuracy: Overall correctness of the model.\n")
    p.add_run("• Precision & Recall: Precision ensures authentic videos aren't falsely flagged, while recall ensures true fakes are reliably detected.\n")
    p.add_run("• ROC-AUC: Demonstrates the model's capability to distinguish between classes at various threshold settings.\n")
    p.add_run("• Confusion Matrix: Visualizes the true positives, true negatives, false positives, and false negatives.")
    
    doc.add_heading("Robustness Testing", level=2)
    add_paragraph("The model was rigorously tested against real-world social media conditions. Deepfake videos were subjected to simulated WhatsApp/Facebook JPEG compression, resizing, and frame-rate alterations. The data augmentation strategy during training successfully fortified the model against severe compression artifacts.")
    doc.add_page_break()

    # 10. Conclusion
    doc.add_heading("10. CONCLUSION", level=1)
    add_paragraph("The project successfully demonstrates the effectiveness of Convolutional Neural Networks, specifically EfficientNet-B4, in addressing the critical challenge of deepfake detection. By employing transfer learning, robust data augmentation, and two-phase training, the system accurately classifies manipulated media.")
    add_paragraph("The fully offline, Flask-based deployment ensures user privacy, eliminating reliance on third-party cloud APIs. Incorporating Grad-CAM explainability transforms the system from a black-box predictor to a transparent forensic tool. This project lays a strong foundation for AI-powered media verification systems, protecting individuals from malicious manipulated media and restoring trust in digital content.")
    doc.add_page_break()

    # 11. Future Scope
    doc.add_heading("11. FUTURE SCOPE", level=1)
    p = doc.add_paragraph()
    p.add_run("• Audio Deepfake Detection: Extending the system to identify voice cloning and speech synthesis using spectrogram-based CNN analysis.\n")
    p.add_run("• Multi-Modal Verification: Combining visual frame analysis with audio analysis to perform lip-sync verification.\n")
    p.add_run("• Temporal Analysis (LSTM): Integrating Long Short-Term Memory (LSTM) networks to analyze temporal inconsistencies between consecutive video frames rather than treating each frame independently.\n")
    p.add_run("• Browser Extension: Deploying the lightweight inference engine as a browser extension to flag deepfakes in real-time on social media platforms.")
    doc.add_page_break()

    # 12. Sample Codes
    doc.add_heading("12. SAMPLE CODES", level=1)
    add_paragraph("Model Definition Snippet (model.py)", bold=True)
    add_paragraph("import torch.nn as nn\nfrom torchvision import models\n\nclass DeepfakeDetector(nn.Module):\n    def __init__(self):\n        super(DeepfakeDetector, self).__init__()\n        self.backbone = models.efficientnet_b4(pretrained=True)\n        num_ftrs = self.backbone.classifier[1].in_features\n        self.backbone.classifier = nn.Identity()\n        self.custom_head = nn.Sequential(\n            nn.Dropout(0.3),\n            nn.Linear(num_ftrs, 512),\n            nn.BatchNorm1d(512),\n            nn.ReLU(),\n            nn.Dropout(0.2),\n            nn.Linear(512, 1),\n            nn.Sigmoid()\n        )\n\n    def forward(self, x):\n        features = self.backbone(x)\n        return self.custom_head(features)", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    # 13. Outcomes
    doc.add_heading("13. OUTCOMES", level=1)
    add_paragraph("The implementation of the system yielded the following key outcomes:")
    p = doc.add_paragraph()
    p.add_run("• Development of a highly accurate CNN model capable of classifying deepfake videos.\n")
    p.add_run("• A fully functional Flask web application allowing drag-and-drop video upload.\n")
    p.add_run("• Instantaneous generation of classification reports including overall verdict, confidence percentage, and frame-by-frame analysis.\n")
    p.add_run("• Successful generation of Grad-CAM heatmaps highlighting manipulated facial regions, greatly improving the forensic utility of the tool.")
    doc.add_page_break()

    # 14. Bibliography
    doc.add_heading("14. BIBLIOGRAPHY", level=1)
    p = doc.add_paragraph()
    p.add_run("1. Rössler, A., et al. \"FaceForensics++: Learning to Detect Manipulated Facial Images.\" arXiv:1901.08971, 2019.\n")
    p.add_run("2. Li, Y., et al. \"Celeb-DF: A Large-scale Challenging Dataset for DeepFake Forensics.\" IEEE CVPR, 2020.\n")
    p.add_run("3. Paszke, A., et al. \"PyTorch: An Imperative Style, High-Performance Deep Learning Library.\" Advances in Neural Information Processing Systems 32, 2019.\n")
    p.add_run("4. Tan, M., and Le, Q. \"EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks.\" ICML, 2019.\n")
    p.add_run("5. Selvaraju, R. R., et al. \"Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization.\" ICCV, 2017.")
    
    # Save the document
    doc.save("e:\\MCA\\MajorProject\\Deepfake_Detection_Using_CNN_Major_Project_Report.docx")

if __name__ == "__main__":
    create_report()
